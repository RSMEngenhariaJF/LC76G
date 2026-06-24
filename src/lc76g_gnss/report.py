"""
lc76g_gnss.report
-----------------
Geração do **relatório do ensaio de precisão** em Word (.docx).

O relatório inclui cabeçalho (título, responsável, data/hora), resumo descritivo,
condições meteorológicas, estatísticas de erro, tabela de pontos (com satélites
ponto a ponto), gráficos (histograma e erro por ponto), um **glossário técnico**
das siglas e um **anexo com os dados brutos** (amostras de cada ponto).

Depende de ``python-docx`` (e, opcionalmente, ``matplotlib`` para os gráficos).
Importe sob demanda: se ``python-docx`` não estiver instalado, a importação deste
módulo falha e o chamador deve tratar isso.
"""

from __future__ import annotations

import os
import tempfile
from typing import List, Optional, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from . import weather as wx
from .accuracy import AccuracyPoint, sample_deviations_m

_AZUL = RGBColor(0x1F, 0x4E, 0x79)
_CINZA = RGBColor(0x59, 0x59, 0x59)


# --------------------------------------------------------------------------- #
# Glossário técnico das siglas/termos que aparecem no relatório
# --------------------------------------------------------------------------- #
GLOSSARY = [
    ("GNSS", "Global Navigation Satellite System — termo genérico para sistemas "
             "de navegação por satélite (GPS, GLONASS, Galileo, BeiDou)."),
    ("Fix", "Solução de posição válida calculada pelo receptor. 2D usa ≥3 "
            "satélites (lat/lon); 3D usa ≥4 (lat/lon/altitude)."),
    ("TTFF", "Time To First Fix — tempo desde a partida até o primeiro fix válido."),
    ("Satélites usados", "Satélites efetivamente empregados no cálculo da "
                         "posição (campo NumSat da sentença GGA)."),
    ("Satélites em vista", "Satélites visíveis/rastreados (sentenças GSV), nem "
                           "todos necessariamente usados na solução."),
    ("NMEA 0183", "Padrão de sentenças de texto emitidas pelo receptor "
                  "(ex.: RMC, GGA, GSA, GSV)."),
    ("RMC", "Recommended Minimum data — posição, velocidade, data/hora e status."),
    ("GGA", "Global Positioning System Fix Data — posição, qualidade do fix, "
            "nº de satélites, HDOP e altitude."),
    ("GSA", "DOP and Active Satellites — tipo de fix (2D/3D) e diluições."),
    ("GSV", "Satellites in View — satélites visíveis por constelação."),
    ("HDOP", "Horizontal Dilution of Precision — fator geométrico do erro "
             "horizontal. Quanto menor, melhor (≤2 é bom em céu aberto)."),
    ("PDOP/VDOP", "Diluições de precisão posicional (3D) e vertical."),
    ("Moda", "Valor mais frequente de uma amostra. Aqui, a posição é estimada "
             "pela moda das latitudes/longitudes (robusta a outliers)."),
    ("Haversine", "Fórmula que calcula a distância sobre a esfera terrestre "
                  "entre duas coordenadas (usada na 'distância medida')."),
    ("Erro (m)", "Diferença entre a distância medida e a informada "
                 "(positivo = superestimou)."),
    ("Erro (%)", "Erro relativo à distância informada (indefinido quando = 0)."),
    ("MAE", "Mean Absolute Error — média dos erros em módulo."),
    ("RMS", "Root Mean Square — raiz da média dos erros ao quadrado; penaliza "
            "erros grandes."),
    ("Desvio padrão", "Dispersão dos erros em torno da média."),
    ("Mediana", "Valor central dos erros ordenados (robusto a outliers)."),
    ("Quartis / IQR", "Q1, mediana e Q3 dividem os dados em quatro partes; o "
                      "IQR (Q3−Q1) mede a dispersão central."),
    ("Diagrama de quartis (boxplot)", "Resumo visual da distribuição: caixa do "
                                      "Q1 ao Q3, traço na mediana, hastes até o "
                                      "alcance e pontos como outliers. Aqui, um "
                                      "por medida, mostra a repetibilidade das "
                                      "amostras de cada ponto."),
    ("Troposfera", "Camada baixa da atmosfera; temperatura, pressão e umidade "
                   "causam atraso troposférico no sinal GNSS."),
    ("Ionosfera", "Camada ionizada da atmosfera; sua atividade atrasa o sinal e "
                  "é a maior fonte de erro em receptores de banda única."),
    ("Índice Kp", "Indicador global (0–9) de perturbação geomagnética/ionosférica; "
                  "valores altos degradam a precisão do GNSS."),
    ("Pressão de superfície", "Pressão atmosférica no local/altitude real."),
    ("Pressão ao nível do mar", "Pressão reduzida ao nível do mar (referência "
                                "meteorológica padrão)."),
    ("UTC", "Tempo Universal Coordenado, emitido pelo GNSS."),
]


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = _AZUL
    return h


def _para(doc, text="", *, bold=False, italic=False, size=11, color=None,
          space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _table(doc, headers, rows, widths=None, header_fill="1F4E79"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(h))
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(cell, header_fill)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run("" if val is None else str(val))
            run.font.size = Pt(8.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    return t


def _fmt(value, suffix="", nd=2):
    if value is None:
        return "—"
    if isinstance(value, float):
        return f"{value:.{nd}f}{suffix}"
    return f"{value}{suffix}"


def _build_charts(points: Sequence[AccuracyPoint], stats: dict, tmpdir: str):
    """Gera PNG (histograma + erro por ponto + boxplot). Retorna caminhos."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        from matplotlib.figure import Figure
    except Exception:
        return []
    errors = [p.error for p in points]
    pcts = [(p.index, p.error_pct) for p in points if p.error_pct is not None]

    paths = []
    fig = Figure(figsize=(8.2, 6.4), dpi=130)
    ax1 = fig.add_subplot(2, 2, 1)
    bins = min(15, max(3, len(errors)))
    ax1.hist(errors, bins=bins, color="#1565c0", edgecolor="white")
    ax1.axvline(stats["mean_error"], color="#c62828", linestyle="--",
                linewidth=1.2, label=f"média {stats['mean_error']:+.2f} m")
    ax1.set_title("Histograma do erro de distância")
    ax1.set_xlabel("Erro (m)")
    ax1.set_ylabel("Frequência")
    ax1.legend(fontsize=8)
    ax1.grid(True, alpha=0.3)

    ax2 = fig.add_subplot(2, 2, 2)
    if pcts:
        idxs = [i for i, _ in pcts]
        ys = [v for _, v in pcts]
        ax2.bar(range(len(ys)), ys, color="#2e7d32", edgecolor="white")
        ax2.set_xticks(range(len(ys)))
        ax2.set_xticklabels([f"P{i}" for i in idxs], fontsize=7)
        ax2.axhline(0, color="#555", linewidth=0.8)
        ax2.set_title("Erro percentual por ponto")
        ax2.set_ylabel("Erro (%)")
    else:
        ab = [p.abs_error for p in points]
        ax2.bar(range(len(ab)), ab, color="#2e7d32", edgecolor="white")
        ax2.set_xticks(range(len(ab)))
        ax2.set_xticklabels([f"P{p.index}" for p in points], fontsize=7)
        ax2.set_title("Erro absoluto por ponto (parado)")
        ax2.set_ylabel("Erro (m)")
    ax2.set_xlabel("Ponto")
    ax2.grid(True, alpha=0.3)

    # Boxplot (diagrama de quartis) da dispersão das amostras de cada ponto.
    ax3 = fig.add_subplot(2, 1, 2)
    devs = [sample_deviations_m(p) for p in points]
    data = [d for d in devs if d]
    labels = [f"P{p.index}" for p, d in zip(points, devs) if d]
    if data:
        ax3.boxplot(data, showmeans=True)
        ax3.set_xticklabels(labels, fontsize=7)
        ax3.set_title("Dispersão das amostras por ponto (quartis)")
        ax3.set_xlabel("Ponto")
        ax3.set_ylabel("Desvio da posição do ponto (m)")
        ax3.grid(True, alpha=0.3)
    else:
        ax3.text(0.5, 0.5, "sem amostras brutas para o boxplot",
                 ha="center", va="center")
        ax3.set_axis_off()
    fig.tight_layout()

    img = os.path.join(tmpdir, "graficos.png")
    fig.savefig(img, dpi=130, bbox_inches="tight")
    paths.append(img)
    return paths


def build_precision_report(path, *, title, responsible, device, mode,
                           reference, points: List[AccuracyPoint], stats: dict,
                           weather=None, sample_count=None, decimals=None,
                           generated_at=None, coord_decimals=6):
    """Gera o relatório .docx do ensaio de precisão em ``path``."""
    def cfmt(v):  # formata coordenada com as casas escolhidas
        return f"{v:.{coord_decimals}f}"

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    modo_txt = ("Por trecho (ponto anterior)" if mode == "trecho"
                else "A partir da origem")
    gerado = generated_at.strftime("%d/%m/%Y %H:%M:%S") if generated_at else "—"
    hora_ini = points[0].time_local if points and points[0].time_local else "—"
    hora_fim = points[-1].time_local if points and points[-1].time_local else "—"

    # ----- Cabeçalho -----
    t = doc.add_heading("Relatório de Ensaio de Precisão GNSS", level=0)
    for r in t.runs:
        r.font.color.rgb = _AZUL
    _para(doc, title or "(sem título)", bold=True, size=14, color=_CINZA,
          space_after=10)

    ref_txt = (f"{cfmt(reference[0])}, {cfmt(reference[1])}" if reference else "—")
    _table(doc, ["Campo", "Conteúdo"], [
        ["Título da medição", title or "—"],
        ["Responsável", responsible or "—"],
        ["Equipamento (DUT)", device or "—"],
        ["Modo de medição", modo_txt],
        ["Início / fim da medição", f"{hora_ini}  →  {hora_fim}"],
        ["Emitido em", gerado],
        ["Posição de origem (lat, lon)", ref_txt],
        ["Nº de pontos", str(len(points))],
        ["Amostras por ponto", str(sample_count) if sample_count else "—"],
        ["Casas decimais da moda", str(decimals) if decimals is not None else "—"],
    ], widths=[6.0, 10.0])
    doc.add_paragraph()

    # ----- 1. Objetivo / metodologia -----
    _heading(doc, "1. Objetivo e metodologia", 1)
    _para(doc,
          "Este relatório apresenta o ensaio de precisão de posicionamento do "
          "receptor GNSS. A posição de referência (origem) e a de cada ponto são "
          "estimadas pela moda de várias amostras de latitude/longitude, "
          "reduzindo o efeito de outliers. A distância medida entre pontos é "
          "calculada pela fórmula de haversine e comparada com a distância "
          "informada pelo operador, gerando o erro (em metros e percentual) e as "
          "estatísticas de desempenho. As siglas estão explicadas no Glossário "
          "(seção 6).")
    _para(doc, f"Modo de medição utilizado: {modo_txt}. "
          + ("Cada ponto é medido contra a origem fixa." if mode != "trecho"
             else "Cada ponto é medido contra o ponto anterior."),
          italic=True, color=_CINZA)

    # ----- 2. Condições meteorológicas -----
    _heading(doc, "2. Condições ambientais (meteorologia e clima espacial)", 1)
    if weather is not None and weather.error is None:
        _para(doc,
              "Condições atmosféricas no momento do ensaio (fonte: Open-Meteo e "
              "NOAA SWPC). Pressão, temperatura e umidade influenciam o atraso "
              "troposférico; o índice Kp indica a perturbação ionosférica — "
              "ambos afetam a precisão do GNSS.")
        _table(doc, ["Variável", "Valor"], [
            ["Temperatura", _fmt(weather.temperature_c, " °C", 1)],
            ["Umidade relativa", _fmt(weather.humidity_pct, " %", 0)],
            ["Pressão de superfície", _fmt(weather.surface_pressure_hpa, " hPa", 1)],
            ["Pressão ao nível do mar", _fmt(weather.pressure_msl_hpa, " hPa", 1)],
            ["Cobertura de nuvens", _fmt(weather.cloud_cover_pct, " %", 0)],
            ["Precipitação", _fmt(weather.precipitation_mm, " mm", 1)],
            ["Vento", _fmt(weather.wind_speed_kmh, " km/h", 1)],
            ["Elevação do terreno", _fmt(weather.elevation_m, " m", 0)],
            ["Índice Kp (ionosfera)", _fmt(weather.kp_index, "", 1)],
            ["Observação (UTC)", weather.observed_utc or "—"],
        ], widths=[7.0, 9.0])
    else:
        motivo = f" ({weather.error})" if weather and weather.error else ""
        _para(doc, "Dados meteorológicos não disponíveis para este ensaio"
              + motivo + ".", italic=True, color=_CINZA)

    # ----- 3. Estatísticas -----
    _heading(doc, "3. Desempenho (estatísticas de erro)", 1)
    if stats.get("n"):
        pct = stats.get("mean_abs_error_pct")
        _para(doc,
              "Resumo das métricas de erro de distância do ensaio. Erro médio "
              "próximo de zero indica ausência de viés; MAE e RMS resumem a "
              "magnitude típica do erro; o desvio padrão indica a dispersão.")
        _table(doc, ["Métrica", "Valor", "Significado"], [
            ["Pontos", str(stats["n"]), "Quantidade de pontos medidos"],
            ["Erro médio", _fmt(stats["mean_error"], " m"), "Viés (com sinal)"],
            ["MAE", _fmt(stats["mean_abs_error"], " m"), "Erro absoluto médio"],
            ["RMS", _fmt(stats["rms_error"], " m"), "Penaliza erros grandes"],
            ["Desvio padrão", _fmt(stats["std_error"], " m"), "Dispersão do erro"],
            ["Mediana", _fmt(stats["median_error"], " m"), "Erro central"],
            ["Erro abs. mínimo", _fmt(stats["min_abs_error"], " m"), "Melhor ponto"],
            ["Erro abs. máximo", _fmt(stats["max_abs_error"], " m"), "Pior ponto"],
            ["Erro % abs. médio",
             _fmt(pct, " %", 1) if pct is not None else "—",
             "Erro relativo médio"],
        ], widths=[4.5, 4.0, 7.5])
    else:
        _para(doc, "Sem pontos medidos.", italic=True, color=_CINZA)

    # ----- 4. Gráficos -----
    with tempfile.TemporaryDirectory() as tmpdir:
        charts = _build_charts(points, stats, tmpdir) if stats.get("n") else []
        if charts:
            _heading(doc, "4. Gráficos", 1)
            for img in charts:
                doc.add_picture(img, width=Cm(16.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ----- 5. Pontos medidos (com satélites ponto a ponto) -----
        _heading(doc, "5. Pontos medidos", 1)
        _para(doc, "Distância informada × medida, erro e satélites (usados na "
              "solução e em vista) em cada ponto. A diferença entre vistos e "
              "usados indica obstrução/qualidade do céu no ponto.")
        rows = []
        for p in points:
            rows.append([
                p.index, _fmt(p.known_distance, "", 2), _fmt(p.measured_distance, "", 2),
                _fmt(p.error, "", 2),
                _fmt(p.error_pct, "", 1) if p.error_pct is not None else "—",
                p.sats if p.sats is not None else "—",
                p.sats_in_view if p.sats_in_view is not None else "—",
                _fmt(p.hdop, "", 2) if p.hdop is not None else "—",
                cfmt(p.lat), cfmt(p.lon),
                (p.time_local or "—"), (p.gnss_utc or "—"),
            ])
        _table(doc, ["#", "Inform. (m)", "Medido (m)", "Erro (m)", "Erro %",
                     "Sat. usados", "Sat. vistos", "HDOP", "Latitude",
                     "Longitude", "Hora local", "UTC GNSS"], rows,
               widths=[0.7, 1.5, 1.5, 1.2, 1.0, 1.3, 1.3, 1.0, 2.2, 2.2, 2.3, 1.6])

        # Documento precisa ser salvo dentro do with se as imagens forem
        # embutidas por referência? Não: add_picture copia os bytes. Mas
        # mantemos o save aqui por simplicidade de fluxo.
        # ----- 6. Glossário -----
        _heading(doc, "6. Glossário técnico", 1)
        _para(doc, "Significado das siglas e termos usados neste relatório.")
        _table(doc, ["Termo", "Descrição"], GLOSSARY, widths=[3.5, 12.5])

        # ----- Anexo A: dados brutos -----
        _heading(doc, "Anexo A — Dados brutos (amostras por ponto)", 1)
        _para(doc, "Amostras de latitude/longitude coletadas em cada ponto, "
              "a partir das quais a posição foi estimada pela moda.",
              italic=True, color=_CINZA)
        any_samples = False
        for p in points:
            if not p.samples:
                continue
            any_samples = True
            view_txt = (f"; vistos {p.sats_in_view}"
                        + (f" [{p.view_breakdown}]" if p.view_breakdown else "")
                        ) if p.sats_in_view is not None else ""
            _para(doc, f"Ponto {p.index} — {len(p.samples)} amostras "
                  f"(informado {p.known_distance:g} m, medido "
                  f"{p.measured_distance:.2f} m; usados "
                  f"{p.sats if p.sats is not None else '—'}{view_txt}):",
                  bold=True, size=10, space_after=2)
            rows = [[i + 1, cfmt(la), cfmt(lo)]
                    for i, (la, lo) in enumerate(p.samples)]
            _table(doc, ["Amostra", "Latitude", "Longitude"], rows,
                   widths=[2.5, 5.0, 5.0])
            doc.add_paragraph()
        if not any_samples:
            _para(doc, "Sem amostras brutas registradas.", italic=True,
                  color=_CINZA)

        doc.save(path)
    return path
